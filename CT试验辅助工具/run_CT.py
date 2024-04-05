import sys
from CT_widget import Ui_Form
from PySide6 import QtWidgets, QtGui, QtCore
from PySide6.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QLineEdit, QLabel, QRadioButton, \
    QTableWidget, QTableWidgetItem
from docx import Document
import os
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
# import win32com.client


# sys.stdout.reconfigure(encoding='utf-8')


class run_CT(Ui_Form, QWidget):
    def __init__(self):
        # super().__init__()
        #
        super().__init__()
        self.setupUi(self)
        self.initUI()
        self.DOCX='1S1-1S2.docx'
        self.input_data = {
            'lineEdit_Imax': '', # 短路电流
            'lineEdit_I1e': '',# 一次额定电流
            'lineEdit_U2': '',# 拐点电压√
            'lineEdit_L_R2': '',# 电阻√
            'lineEdit_I2e': '',# 二次额定电流√
            'lineEdit_Zreal': '',# 实际二次负担
            'lineEdit_K': '',# 可靠系数
            'lineEdit_K1': '',# 可靠系数1
            'lineEdit_K2': '',# 可靠系数2
            'key10': ''
        }

    def initUI(self):
        self.bind()

    def bind(self):
        # self.btn_culc.clicked.connect(lambda: self.btn_slot('trans_main'))  # 变压器主保护
        self.btn_culc.clicked.connect(self.culc)
        self.btn_save2doc.clicked.connect(self.save2doc)

        self.lineEdit_Imax.textChanged.connect(lambda: self.save_data('lineEdit_Imax'))
        self.lineEdit_I1e.textChanged.connect(lambda: self.save_data('lineEdit_I1e'))
        self.lineEdit_I2e.textChanged.connect(lambda: self.save_data('lineEdit_I2e'))

        self.lineEdit_U2.textChanged.connect(lambda: self.save_data('lineEdit_U2'))
        self.lineEdit_L_R2.textChanged.connect(lambda: self.save_data('lineEdit_L_R2'))
        self.lineEdit_Zreal.textChanged.connect(lambda: self.save_data('lineEdit_Zreal'))
        self.lineEdit_K.textChanged.connect(lambda: self.save_data('lineEdit_K'))
        self.lineEdit_K1.textChanged.connect(lambda: self.save_data('lineEdit_K1'))
        self.lineEdit_K2.textChanged.connect(lambda: self.save_data('lineEdit_K2'))

    def culc(self):
        print(self.input_data)
        self.culc_Z2 = float(self.input_data['lineEdit_L_R2']) * float(self.input_data['lineEdit_K1'])  # K1*R2
        self.culc_m = float(self.input_data['lineEdit_K']) * float(self.input_data['lineEdit_Imax']) / float(
            self.input_data['lineEdit_I1e'])  # m = K * Imax / I1e
        if self.input_data['lineEdit_I2e'] == '5':
            self.culc_I0 = self.culc_m / 2
        elif self.input_data['lineEdit_I2e'] == '1':
            self.culc_I0 = self.culc_m / 10
        else:
            self.culc_I0 = self.culc_m / 2
        self.culc_E = float(self.input_data['lineEdit_U2']) - self.culc_I0 * self.culc_Z2  # E = U2 - I0 * Z2

        self.culc_Zallow = float(self.input_data['lineEdit_U2']) / (9 * self.culc_I0) - self.culc_Z2 * float(
            self.input_data['lineEdit_K2'])  # 允许负担 Z允 = U2 / 9Io - K2 * Z2
        if self.culc_Zallow > float(self.input_data['lineEdit_Zreal']):
            self.is_OK = True
        else:
            self.is_OK = False

        print(self.culc_E, self.culc_m, self.culc_Z2, self.culc_I0)

        measure_culc_Z2_item = QTableWidgetItem(str(self.culc_Z2))
        self.tableWidget.setItem(0, 0, measure_culc_Z2_item)

        measure_culc_m_item = QTableWidgetItem(str(self.culc_m))
        self.tableWidget.setItem(0, 1, measure_culc_m_item)
        measure_culc_I0_item = QTableWidgetItem(str(self.culc_I0))
        self.tableWidget.setItem(0, 2, measure_culc_I0_item)

        self.culc_E = round(self.culc_E, 3)
        measure_culc_E_item = QTableWidgetItem(str(self.culc_E))
        self.tableWidget.setItem(0, 3, measure_culc_E_item)
        self.culc_Zallow = round(self.culc_Zallow, 3)
        measure_culc_Zallow_item = QTableWidgetItem(str(self.culc_Zallow))

        self.tableWidget.setItem(0, 4, measure_culc_Zallow_item)

        measure_is_OK_item = QTableWidgetItem(str(self.is_OK))
        self.tableWidget.setItem(0, 5, measure_is_OK_item)

    def save2doc(self):
        # word_app = win32com.client.Dispatch("Word.Application")



        # 读取原始文档
        doc_input = Document(self.DOCX)
        for paragraph in doc_input.paragraphs:
            # 检查段落文本是否包含“拐点电压Vkn”
            if '拐点电压Vkn' in paragraph.text:
                # 找到包含“拐点电压Vkn”的段落，提取并打印值
                value = paragraph.text.split('Vkn:')[1].split(' V')[0]
                print(f"拐点电压Vkn的值为: {value}")
                break  # 找到后就可以停止搜索
        print('1')
        # 创建一个新的Word文档
        doc_output = Document()

        # 用于存储励磁取整数据的列表
        excitation_data = []

        # 遍历文档中的表格，找到励磁取整数据
        for table in doc_input.tables:
            for row in table.rows:
                if '励磁取整数据' in str(row):
                    # 提取数据
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            excitation_data.append(text)

        # 检查是否找到励磁取整数据
        if not excitation_data:
            print("没有找到励磁取整数据，请检查文档内容。")
        else:
            # 将数据写入新的Word文档
            doc_output.add_heading('励磁取整数据', 0)

            # 计算列数
            col_count = len(excitation_data[0].split('\n')[0].split('\n')) + 1

            # 添加表格
            table = doc_output.add_table(rows=1, cols=col_count)
            header_cells = table.rows[0].cells
            for header in excitation_data[0].split('\n')[0].split('\n'):
                header_cells[0].text = '序号'
                header_cells[1].text = header
                excitation_data.pop(0)

            # 添加数据行
            for i, row_data in enumerate(excitation_data):
                if row_data:  # 确保行数据不为空
                    row = table.add_row()
                    cells = row.cells
                    for j, cell_data in enumerate(row_data.split('\n')):
                        if cell_data:  # 确保单元格数据不为空
                            cells[j].text = cell_data

            # 保存新的Word文档
            doc_output.save('path_to_your_output_document.docx')

    def find_table_with_target_row(self, doc, target_row):
        for para in doc.paragraphs:
            # 查找包含“CT试验参数:”的段落
            if target_row in para.text:
                # 获取该段落的下一段落的表格
                table = para.next_paragraph.tables[0]
                # 遍历表格的每一行
                for row in table.rows:
                    # 查找包含“额定二次电流Isn”的单元格
                    if '额定二次电流Isn' in row.cells[0].text:
                        # 获取额定电流的值
                        rated_current = row.cells[1].text
                        print(rated_current)
                        break

    def save_data(self, tag):
        if tag == 'lineEdit_Imax':
            self.input_data['lineEdit_Imax'] = self.lineEdit_Imax.text()

        # 变压器后备
        elif tag == 'lineEdit_I1e':
            self.input_data['lineEdit_I1e'] = self.lineEdit_I1e.text()
        elif tag == 'lineEdit_I2e':
            self.input_data['lineEdit_I2e'] = self.lineEdit_I2e.text()

        elif tag == 'lineEdit_U2':
            self.input_data['lineEdit_U2'] = self.lineEdit_U2.text()

        elif tag == 'lineEdit_L_R2':
            self.input_data['lineEdit_L_R2'] = self.lineEdit_L_R2.text()

        elif tag == 'lineEdit_Zreal':
            self.input_data['lineEdit_Zreal'] = self.lineEdit_Zreal.text()

        elif tag == 'lineEdit_K':
            self.input_data['lineEdit_K'] = self.lineEdit_K.text()

        elif tag == 'lineEdit_K1':
            self.input_data['lineEdit_K1'] = self.lineEdit_K1.text()

        elif tag == 'lineEdit_K2':
            self.input_data['lineEdit_K2'] = self.lineEdit_K2.text()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = run_CT()
    ex.show()
    sys.exit(app.exec_())
